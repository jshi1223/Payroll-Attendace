import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Theme ──
RED = RGBColor(0xCC, 0x00, 0x00)
DARK_RED = RGBColor(0x8B, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)

IMGDIR = os.path.join(os.path.dirname(__file__), 'images')
OUTPUT = os.path.join(os.path.dirname(__file__), 'KVSK_Payroll_Attendance_System_Manual.docx')

doc = Document()

# ── Default ──
sty = doc.styles['Normal']
sty.font.name = 'Calibri'
sty.font.size = Pt(10.5)
sty.font.color.rgb = DARK_GRAY
sty.paragraph_format.space_after = Pt(6)
sty.paragraph_format.line_spacing = 1.15

# ── Heading 1 (red bottom border) ──
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(18)
h1.font.color.rgb = RED
h1.font.bold = True
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(10)
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="12" w:space="4" w:color="CC0000"/>'
    f'</w:pBdr>'
)
h1.element.get_or_add_pPr().append(pBdr)

# ── Heading 2 ──
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(14)
h2.font.color.rgb = BLACK
h2.font.bold = True
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(6)

# ── Caption ──
cap = doc.styles['Caption']
cap.font.name = 'Calibri'
cap.font.size = Pt(9)
cap.font.italic = True
cap.font.color.rgb = GRAY
cap.paragraph_format.space_before = Pt(2)
cap.paragraph_format.space_after = Pt(14)
cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Page setup ──
for sec in doc.sections:
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

# ── Helpers ──

def shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    )

def add_img(name, caption_text, width=Inches(5.2)):
    path = os.path.join(IMGDIR, name)
    if not os.path.exists(path):
        p = doc.add_paragraph(f'[Image: {name}]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    run.add_picture(path, width=width)
    doc.add_paragraph(caption_text, style='Caption')

def para(en, fil=None):
    doc.add_paragraph(en)
    if fil:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(fil)
        r.font.color.rgb = GRAY
        r.italic = True

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.add_run(text)

def step(num, en, fil):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'{num}. ')
    r.bold = True
    r.font.color.rgb = RED
    p.add_run(en)
    if fil:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Cm(1)
        r2 = p2.add_run(fil)
        r2.font.color.rgb = GRAY
        r2.italic = True

def pb():
    doc.add_page_break()

def red_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, 'CC0000')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_GRAY
            if ri % 2 == 1:
                shade(c, 'F5F5F5')
    return t

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f'\u2139 {text}')
    r.font.color.rgb = DARK_RED
    r.font.size = Pt(9.5)
    r.italic = True


# ═══════════════════════════
#  COVER PAGE
# ═══════════════════════════
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('_' * 60)
r.font.color.rgb = RED
r.font.size = Pt(10)

doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('KVSK Payroll\nAttendance System')
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = BLACK

doc.add_paragraph()

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('User Manual')
r.font.size = Pt(14)
r.font.color.rgb = RED

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('_' * 60)
r.font.color.rgb = RED
r.font.size = Pt(10)

doc.add_paragraph()

v = doc.add_paragraph()
v.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = v.add_run('Version 1.0  |  July 2026')
r.font.size = Pt(11)
r.font.color.rgb = GRAY

pb()

# ═══════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════
doc.add_heading('Table of Contents', level=1)
toc = [
    '1.  System Overview',
    '2.  Login and Logout',
    '3.  Dashboard',
    '4.  Payroll Processing',
    '5.  Attendance Management',
    '6.  Employee Management',
    '7.  Cash Advances (Bale)',
    '8.  Archive',
    '9.  Audit Trail',
    '10. Dark Mode',
    '11. Payslip',
    '12. Keyboard Shortcuts',
    '13. Change Password',
    '14. Bulk Print Payslips',
    '15. Mobile View',
    '16. Bale System (Cash Advance Flow)',
    '17. Payroll Computation (as shown on Payslip)',
    '18. Lock and Unlock Flow',
    '19. Session and Security',
    '20. Troubleshooting and FAQ',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(10.5)
    p.runs[0].font.color.rgb = DARK_GRAY

pb()

# ═══════════════════════════
#  1. SYSTEM OVERVIEW
# ═══════════════════════════
doc.add_heading('1. System Overview', level=1)
para(
    'The KVSK Payroll Attendance System is a web-based application designed for managing employee attendance, '
    'processing weekly payroll, handling cash advances (bale), and generating payslips. It supports two user '
    'roles: Admin (full access) and User (limited access). The system uses a PostgreSQL database and runs on '
    'a Node.js backend with an HTML/CSS/JavaScript frontend.',
    'Ang KVSK Payroll Attendance System ay isang web-based na aplikasyon para sa pamamahala ng attendance ng '
    'empleyado, pagproseso ng weekly payroll, pangangasiwa ng cash advances (bale), at paggawa ng payslip. Sumusuporta ito '
    'sa dalawang user roles: Admin (buong access) at User (limitadong access).'
)

doc.add_heading('Navigation', level=2)
para(
    'The sidebar on the left provides access to the main modules: Dashboard, Payroll, Attendance, Employees, '
    'and Archive. Each module corresponds to a numbered shortcut key (1\u20135).',
    'Ang sidebar sa kaliwa ay nagbibigay ng access sa mga pangunahing modules: Dashboard, Payroll, Attendance, '
    'Employees, at Archive. Bawat module ay may katumbas na shortcut key (1\u20135).'
)

doc.add_heading('User Roles', level=2)
red_table(
    ['Role', 'Access', 'Responsibilities'],
    [
        ['Admin', 'Full system access',
         'Manage all modules, audit trail, unlock payroll, delete records'],
        ['User', 'Limited access',
         'View and manage attendance, record payments, limited payroll actions'],
    ]
)

pb()

# ═══════════════════════════
#  2. LOGIN AND LOGOUT
# ═══════════════════════════
doc.add_heading('2. Login and Logout', level=1)

doc.add_heading('2.1 Logging In', level=2)
step('1', 'Open your web browser and navigate to the system URL.',
     'Buksan ang iyong web browser at pumunta sa URL ng sistema.')
step('2', 'Enter your username and password in the login form.',
     'Ilagay ang iyong username at password sa login form.')
step('3', 'Click the Login button to access the dashboard.',
     'I-click ang Login button upang makapasok sa dashboard.')

add_img('screenshot-01-login.png', 'Figure 2.1: Login Screen', Inches(4.0))
add_img('screenshot-01b-login-filled.png', 'Figure 2.2: Login Form with Credentials', Inches(4.0))

doc.add_heading('2.2 Logging Out', level=2)
para(
    'Click your username at the bottom of the sidebar, then select "Logout" from the menu. '
    'You can also press Ctrl+Shift+Q on your keyboard to log out instantly.',
    'I-click ang iyong username sa ibabang bahagi ng sidebar, pagkatapos ay piliin ang "Logout" mula sa menu. '
    'Bilang alternatibo, pindutin ang Ctrl+Shift+Q para agad na mag-logout.')

doc.add_heading('2.3 Session Timeout', level=2)
para(
    'Sessions expire after 8 hours of inactivity. You will be redirected to the login page automatically when '
    'your session expires.',
    'Ang session ay mag-e-expire pagkatapos ng 8 oras na walang aktibidad. Awtomatiko kang ire-redirect sa login page.')

pb()

# ═══════════════════════════
#  3. DASHBOARD
# ═══════════════════════════
doc.add_heading('3. Dashboard', level=1)
para(
    'The Dashboard is the main landing page after login. It displays summary cards with key metrics '
    'and two "Needs Attention" sections.',
    'Ang Dashboard ay ang pangunahing pahina pagkatapos mag-login.')

doc.add_heading('Summary Cards', level=2)
para('Six cards show: total employees, today\'s attendance count, total paid amount, remaining salary balance, '
     'bale balance, and previous unpaid carryover.')

doc.add_heading('Needs Attention', level=2)
bullet('No Attendance Today \u2014 Lists employees who have not clocked in today. Shows days since last attendance. '
       'Has a "Mark All Present" button and a "Take Attendance" link.')
bullet('Unpaid This Period \u2014 Lists employees with unpaid salary for the current week. '
       'Has a "Go to Payroll" link.')

add_img('screenshot-02-dashboard.png', 'Figure 3.1: Dashboard Overview')

pb()

# ═══════════════════════════
#  4. PAYROLL PROCESSING
# ═══════════════════════════
doc.add_heading('4. Payroll Processing', level=1)
para(
    'The Payroll view (Ctrl+2) shows all active employees with their weekly payroll data. '
    'Use the date toolbar at the top to switch between pay periods (Weekly or Semi-Monthly). '
    'The table includes employee number, name, daily rate, days worked, gross salary, previous balances, '
    'cash advances, extra payments, paid amount, current balance, bale balance, and payment status.',
    'Ang Payroll view ay nagpapakita ng lahat ng active na empleyado kasama ang kanilang weekly payroll data.')

add_img('screenshot-03-payroll.png', 'Figure 4.1: Payroll Records Table')

doc.add_heading('4.1 Managing a Payroll Entry', level=2)
para(
    'Click the "Manage" button next to an employee to open the Payroll Entry Modal. From here you can:')
bullet('Record salary payments ("Bayad Sahod")')
bullet('Add cash advances ("Bale / Utang")')
bullet('Add extra payments ("Dagdag Sahod")')
bullet('Record bale repayments ("Bayad Bale")')
bullet('Preview the payslip ("Tingnan Payslip")')
bullet('Generate the official payslip ("Generate Payslip")')

add_img('screenshot-04-payroll-entry-modal.png', 'Figure 4.2: Payroll Entry Modal')

doc.add_heading('4.2 Adding Transactions', level=2)
para(
    'Inside the Payroll Entry Modal, each section has an "Add" button. Click it to open a transaction form '
    'where you enter the amount, select a date, and add notes. Click "Save" to record.',
    'Sa loob ng Payroll Entry Modal, bawat section ay may "Add" button. I-click ito para buksan ang transaction form.')

add_img('screenshot-05-add-transaction.png', 'Figure 4.3: Add Transaction Form', Inches(4.5))

doc.add_heading('4.3 Payslip Preview', level=2)
para(
    'Click "Tingnan Payslip" at the bottom of the modal to see the full payslip breakdown. From here you can '
    'print or download as PDF, or click "Generate Payslip" to finalize and lock the entry.',
    'I-click ang "Tingnan Payslip" sa ibaba ng modal upang makita ang buong payslip breakdown.')

add_img('screenshot-13-payslip.png', 'Figure 4.4: Payslip Preview', Inches(4.5))

pb()

# ═══════════════════════════
#  5. ATTENDANCE
# ═══════════════════════════
doc.add_heading('5. Attendance Management', level=1)
para(
    'The Attendance view (Ctrl+3) records daily attendance. Select a date, then click "Add Attendance" '
    'to log an employee. The table shows employee number, name, time-in, time-out, daily rate, and daily salary. '
    'You can search by employee name or filter by employee. Export to CSV using the Export button.',
    'Ang Attendance view ay ginagamit para i-record ang daily attendance.')

add_img('screenshot-06-attendance.png', 'Figure 5.1: Attendance View')

pb()

# ═══════════════════════════
#  6. EMPLOYEES
# ═══════════════════════════
doc.add_heading('6. Employee Management', level=1)
para(
    'The Employees view (Ctrl+4) manages all employee records. Each record includes: employee number, name, '
    'daily rate, phone number, SSS/PhilHealth/Pag-IBIG/TIN numbers, and pay period (7, 14, 21, or 30 days). '
    'Search across any field using the search bar.',
    'Ang Employees view ay nagbibigay-daan upang pamahalaan ang lahat ng record ng empleyado.')

add_img('screenshot-07-employees.png', 'Figure 6.1: Employee List')

doc.add_heading('6.1 Adding an Employee', level=2)
para(
    'Click "Add Employee" above the table to open the modal. Fill in all required fields and click "Save". '
    'The employee number is auto-generated as EMP-xxxxx.')
add_img('screenshot-08-employee-modal.png', 'Figure 6.2: Add Employee Modal', Inches(4.5))

doc.add_heading('6.2 Editing an Employee', level=2)
para(
    'Click the edit icon (pencil) next to any employee row. Modify the fields and click "Save" to update.')
add_img('screenshot-08b-edit-employee.png', 'Figure 6.3: Edit Employee Modal', Inches(4.5))

pb()

# ═══════════════════════════
#  7. CASH ADVANCES
# ═══════════════════════════
doc.add_heading('7. Cash Advances (Bale)', level=1)
para(
    'Cash advances can be managed through the Payroll Entry Modal. From the payroll table, click "Manage" '
    'on any employee, then select "Bale / Utang" and click "Add" to record a new cash advance. '
    'The system allows one cash advance per employee per day.',
    'Ang cash advances ay maaaring pamahalaan sa pamamagitan ng Payroll Entry Modal.')

add_img('screenshot-09-cash-advances.png', 'Figure 7.1: Cash Advances List')

pb()

# ═══════════════════════════
#  8. ARCHIVE
# ═══════════════════════════
doc.add_heading('8. Archive', level=1)
para(
    'The Archive view (Ctrl+6) contains inactive employees. Restore an employee by clicking "Restore" to '
    'make them active again. Archived employees are excluded from attendance, payroll, and other operations.',
    'Ang Archive view ay naglalaman ng mga inactive na empleyado.')

add_img('screenshot-10-archive.png', 'Figure 8.1: Archive View')

pb()

# ═══════════════════════════
#  9. AUDIT TRAIL
# ═══════════════════════════
doc.add_heading('9. Audit Trail', level=1)
para(
    'The Audit Trail (admin only, Ctrl+Shift+A) logs all system actions: logins, employee changes, payroll '
    'generation and unlock, payments, and more. You can filter by entity, action type, date range, and search text. '
    'Export results to CSV using the Export button.',
    'Ang Audit Trail ay nagla-log ng lahat ng aksyon sa sistema. Admin users lang ang may access.')

add_img('screenshot-11-audit-trail.png', 'Figure 9.1: Audit Trail Modal', Inches(5.5))

pb()

# ═══════════════════════════
#  10. DARK MODE
# ═══════════════════════════
doc.add_heading('10. Dark Mode', level=1)
para(
    'Toggle dark mode using the moon/sun icon in the sidebar or the mobile top bar. '
    'Your preference is saved and persists across sessions.',
    'I-toggle ang dark mode gamit ang moon/sun icon sa sidebar.')

add_img('screenshot-12-dark-mode.png', 'Figure 10.1: Dark Mode View')

pb()

# ═══════════════════════════
#  11. KEYBOARD SHORTCUTS
# ═══════════════════════════
doc.add_heading('11. Keyboard Shortcuts', level=1)
para(
    'Press the ? key at any time to open the shortcuts help modal. Available shortcuts:',
    'Pindutin ang ? key anumang oras upang makita ang shortcuts.')

red_table(
    ['Key', 'Action'],
    [
        ['1\u20135', 'Switch between Dashboard (1), Payroll (2), Attendance (3),\nEmployees (4), Archive (5)'],
        ['/', 'Focus the search bar'],
        ['\u2190 / \u2192', 'Previous / Next period in Payroll, Dashboard,\nand Attendance views'],
        ['?', 'Open this shortcuts help modal'],
        ['Esc', 'Close modal, blur input, or cancel'],
        ['Ctrl+Shift+Q', 'Logout'],
        ['Ctrl+Shift+A', 'Open Audit Trail (admin only)'],
    ])

add_img('screenshot-14-shortcuts.png', 'Figure 11.1: Keyboard Shortcuts Help', Inches(5.0))

pb()

# ═══════════════════════════
#  12. CHANGE PASSWORD
# ═══════════════════════════
doc.add_heading('12. Change Password', level=1)
para(
    'Click your username at the bottom of the sidebar, then select "Change Password". '
    'Enter your current password, then your new password twice for confirmation. Click "Save" to apply. '
    'The change takes effect immediately.',
    'I-click ang iyong username sa sidebar at piliin ang "Change Password". Ilagay ang iyong kasalukuyang '
    'at bagong password.')

add_img('screenshot-15-change-password.png', 'Figure 12.1: Change Password Screen')
note('The default admin password is configured in the .env file (BOOTSTRAP_PASSWORD). Change it immediately after first login.')

pb()

# ═══════════════════════════
#  13. BULK PRINT
# ═══════════════════════════
doc.add_heading('13. Bulk Print Payslips', level=1)
para(
    'The "Bulk Print Payslips" button in the Payroll view toolbar prints all generated payslips for the '
    'current week at once. The button shows a counter (e.g., "Bulk Print (5/8)") indicating how many payslips '
    'have been generated. All payslips must be generated before bulk printing is enabled.',
    'Ang "Bulk Print Payslips" button ay nagpi-print ng lahat ng generated payslip sa isang batch.')

pb()

# ═══════════════════════════
#  14. MOBILE VIEW
# ═══════════════════════════
doc.add_heading('14. Mobile View', level=1)
para(
    'The system is responsive and works on mobile devices. On screens narrower than 768px:')
bullet('The sidebar is hidden; tap the hamburger (\u2630) icon to open it as an overlay.')
bullet('A bottom navigation bar appears with the same five views.')
bullet('Tables scroll horizontally with sticky first columns.')
bullet('All major features are accessible.')

add_img('screenshot-16-mobile-dashboard.png', 'Figure 14.1: Mobile Dashboard', Inches(3.5))
add_img('screenshot-17-mobile-payroll.png', 'Figure 14.2: Mobile Payroll at 600px width', Inches(4.8))
add_img('screenshot-18-mobile-sidebar.png', 'Figure 14.3: Mobile Sidebar Overlay', Inches(3.5))

doc.add_heading('14.1 Tablet View', level=2)
para(
    'On tablets (768px and above), the sidebar is hidden by default but can be opened via the hamburger icon. '
    'The Payroll table is more readable with most columns visible and horizontal scroll for the rest.',
    'Sa tablet, ang sidebar ay nakatago pero maaaring buksan sa pamamagitan ng hamburger icon.')

add_img('screenshot-19-tablet-payroll.png', 'Figure 14.4: Payroll at Tablet Width (768px)', Inches(5.5))

pb()

# ═══════════════════════════
#  15. BALE SYSTEM
# ═══════════════════════════
doc.add_heading('15. Bale System (Cash Advance Flow)', level=1)
para(
    'The bale system lets employees take cash advances against future salary. '
    'The system tracks two separate balances: salary balance (unpaid salary) and bale balance '
    '(outstanding advances).')

doc.add_heading('How It Works', level=2)
bullet('Taking an advance: In the Payroll Entry Modal, go to "Bale / Utang" and click Add. '
        'One advance per employee per day.')
bullet('Repaying: In the same modal, go to "Bayad Bale" and record the payment. '
        'The payment reduces the bale balance.')
bullet('Carryover: Bale balances carry over week to week until fully repaid.')
bullet('Display: The Payroll table shows "Total Adv." (current week advances) and "Adv. Bal." (running bale balance).')

doc.add_heading('Bale Deduction Option', level=2)
para(
    'When processing payroll, you can check the "Bale Deduction" option to automatically apply '
    'the bale balance as a deduction from the employee\'s pay.')

pb()

# ═══════════════════════════
#  16. PAYROLL COMPUTATION
# ═══════════════════════════
doc.add_heading('16. Payroll Computation (as shown on Payslip)', level=1)
para(
    'The payslip displays a complete breakdown of earnings and deductions. Below is the exact layout '
    'matching the printed payslip.',
    'Ang payslip ay nagpapakita ng kumpletong breakdown ng kita at deductions.')

doc.add_heading('Table A \u2014 Daily Breakdown', level=2)
para('For each day in the pay period, the payslip lists: the date, whether the employee came in, '
     'the daily earnings, any bale taken that day, and the running total.')

red_table(
    ['Column', 'Meaning'],
    [
        ['Araw', 'Date of the work day'],
        ['Pumasok', 'Checked if the employee came in (1.00) or not (\u2013)'],
        ['Kita', 'Daily earnings (rate \u00d7 1 if present)'],
        ['Bale (Utang)', 'Cash advance taken that day, if any'],
        ['Kabuuan', 'Running total of earnings for the period'],
    ])

doc.add_heading('Table B \u2014 Kabuuang Kita (Total Earnings)', level=2)
red_table(
    ['Line Item', 'How It Is Computed'],
    [
        ['Sahod (rate \u00d7 days)', 'Days present \u00d7 daily rate'],
        ['Dagdag Sahod', 'Extra payments / bonuses added this period'],
        ['Natitira mula sa Nakaraan', 'Unpaid salary carried over from previous periods'],
        ['Kabuuang Kita', 'Sahod + Dagdag Sahod + Previous Unpaid'],
        ['Bayad Bale (Utang)', 'Bale repayment amount (deduction)'],
    ])

doc.add_heading('Table C \u2014 Sahod na Tatangapin (Take-Home Pay)', level=2)
red_table(
    ['Line Item', 'How It Is Computed'],
    [
        ['SAHOD NA TATANGGAPIN', 'Kabuuang Kita \u2212 Bayad Bale'],
        ['Bayad Sahod', 'Salary payment made this period'],
        ['BALANSA', 'Take-home pay \u2212 Bayad Sahod (remaining unpaid)'],
    ])

doc.add_heading('Table D \u2014 Balensa ng Bale (Bale Balance)', level=2)
red_table(
    ['Line Item', 'How It Is Computed'],
    [
        ['Utang mula sa Nakaraan', 'Bale balance carried over from previous period'],
        ['Bale ngayong Period', 'New cash advances taken this period'],
        ['Kabuuang Utang (Bale)', 'Previous bale + new advances'],
        ['Bayad Bale ngayong Period', 'Bale repayment this period'],
        ['Natitirang Utang (Bale)', 'Total bale \u2212 bale payment (remaining)'],
    ])

doc.add_paragraph()
para(
    'Payment can also go first to any previous unpaid balance. The "Balance" column on the Payroll table '
    'shows remaining unpaid salary, while "Adv. Bal." shows remaining bale to repay. These are tracked separately.',
    'Ang "Balance" column ay nagpapakita ng natitira pang hindi nababayarang sahod. Ang "Adv. Bal." ay '
    'natitira pang bale na kailangang bayaran.')

pb()

# ═══════════════════════════
#  17. LOCK / UNLOCK
# ═══════════════════════════
doc.add_heading('17. Lock and Unlock Flow', level=1)
para(
    'Once a payslip is generated, the payroll entry becomes "Locked" to prevent further changes. '
    'This ensures data integrity after payroll finalization.',
    'Kapag na-generate ang payslip, ang payroll entry ay nagiging "Locked" upang hindi na mabago.')

doc.add_heading('Lock States', level=2)
red_table(
    ['State', 'What It Means'],
    [
        ['Unlocked', 'Entry can be modified. "Manage" button is shown.'],
        ['Generated (Locked)', 'Payslip finalized. Shows "Locked" badge. Admin can "Unlock".'],
        ['Locked (Same Period)', 'Entry locked with matching period. Cannot be modified.'],
    ])

doc.add_heading('Unlocking', level=2)
step('1', 'Find the locked employee in the Payroll table.',
     'Hanapin ang naka-lock na empleyado sa Payroll table.')
step('2', 'Click the "Unlock" button next to their row.',
     'I-click ang "Unlock" button.')
step('3', 'Make your changes.',
     'Gawin ang mga pagbabago.')
step('4', 'Re-generate the payslip.',
     'I-generate muli ang payslip.')

pb()

# ═══════════════════════════
#  18. SESSION & SECURITY
# ═══════════════════════════
doc.add_heading('18. Session and Security', level=1)

doc.add_heading('18.1 Session Management', level=2)
red_table(
    ['Feature', 'Details'],
    [
        ['Storage', 'Server-side sessions via HTTP-only cookies'],
        ['Duration', '8 hours of inactivity before automatic logout'],
        ['Concurrency', 'One active session per user'],
    ])

doc.add_paragraph()
doc.add_heading('18.2 Security Measures', level=2)
bullet('CSRF Protection \u2014 All modifying requests require a valid CSRF token.')
bullet('Rate Limiting \u2014 Max 10 login attempts per 15 minutes per IP.')
bullet('Audit Logging \u2014 All financial and admin actions logged with timestamps.')
bullet('Password Hashing \u2014 bcrypt with cost factor 12.')
bullet('Security Headers \u2014 CSP, HSTS, X-Frame-Options, and others enforced.')

pb()

# ═══════════════════════════
#  19. TROUBLESHOOTING / FAQ
# ═══════════════════════════
doc.add_heading('19. Troubleshooting and FAQ', level=1)

faqs = [
    ('I forgot my password.',
     'Contact your system administrator to reset it.'),
    ('Cannot log in.',
     'Check your username and password. After 10 failed attempts, you will be locked out for 15 minutes.'),
    ('How to generate a payslip?',
     'Go to Payroll, click "Manage" on the employee, review data, then click "Generate Payslip".'),
    ('What does "Locked" mean?',
     'The payslip has been finalized. Click "Unlock" to make changes, then re-generate.'),
    ('How to undo a payment?',
     'Admin users can delete a payment from the Payroll Entry Modal. Recorded in the audit trail.'),
    ('How to print multiple payslips?',
     'Use "Bulk Print Payslips" in the Payroll view after generating all payslips.'),
    ('Default password?',
     'Set in the .env file (BOOTSTRAP_PASSWORD). Change immediately after first login.'),
    ('System not loading?',
     'Refresh (F5). If the problem persists, check that the server and database are running.'),
]

for q, en in faqs:
    p = doc.add_paragraph()
    r = p.add_run(f'Q: {q}')
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RED
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph(en)

# ═══ SAVE ═══
doc.save(OUTPUT)
print(f'Manual generated: {OUTPUT}')
print(f'Size: {os.path.getsize(OUTPUT) / 1024:.0f} KB')
